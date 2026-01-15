import streamlit as st
import sys
import subprocess
import time
import os
import pandas as pd  # Essential for dataframes
import urllib.parse
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, ForeignKey, Text
from sqlalchemy.orm import sessionmaker, declarative_base, relationship
from werkzeug.security import generate_password_hash, check_password_hash
from io import BytesIO

# --- PAGE CONFIGURATION ---
st.set_page_config(page_title="GNC Admin Dashboard", page_icon="🎓", layout="wide")

if "page" not in st.session_state:
    st.session_state.page = "HOME"

if "selected_department_id" not in st.session_state:
    st.session_state.selected_department_id = None

if "selected_subject_id" not in st.session_state:
    st.session_state.selected_subject_id = None

if "generated_paper" not in st.session_state:
    st.session_state.generated_paper = None


# --- COLLEGE / AUTH CONFIG ---
COLLEGE_NAME = "Guru Nanak College"
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")

# --- CUSTOM CSS ---
st.markdown("""
<style>
    .dept-card {
        background-color: #262730; color: white; padding: 20px;
        border-radius: 8px; text-align: center; height: 180px;
        display: flex; flex-direction: column; align-items: center; justify-content: center;
        margin-bottom: 15px; border: 1px solid #3d3d3d;
        box-shadow: 0 4px 6px rgba(0,0,0,0.3); transition: transform 0.2s;
    }
    .dept-card:hover { transform: scale(1.02); border-color: #ff4b4b; }
    .dept-name { font-size: 18px; font-weight: 600; margin-bottom: 12px; }
    .dept-badge { background-color: #444; color: #ddd; padding: 4px 10px; border-radius: 12px; font-size: 12px; }
</style>
""", unsafe_allow_html=True)

# --- DATABASE CONNECTION ---
db_user = "root"
raw_password = "abiR@3121" 
encoded_password = urllib.parse.quote_plus(raw_password)
db_host = "localhost"
db_name = "qp_generator"
DB_URL = f'mysql+pymysql://{db_user}:{encoded_password}@{db_host}/{db_name}'
UPLOAD_FOLDER = 'uploads'

try:
    engine = create_engine(
        DB_URL,
        pool_size=10, max_overflow=20, pool_pre_ping=True, pool_recycle=3600, echo=False
    )
    with engine.connect() as conn: pass
except Exception as e:
    st.error(f"❌ Database Connection Error: {e}")
    st.stop()

Session = sessionmaker(bind=engine)
Base = declarative_base()

# --- MODELS ---
class School(Base):
    __tablename__ = 'school'
    id = Column(Integer, primary_key=True)
    name = Column(String(100))
    departments = relationship("Department", back_populates="school", cascade="all, delete-orphan")

class Department(Base):
    __tablename__ = 'department'
    id = Column(Integer, primary_key=True)
    school_id = Column(Integer, ForeignKey('school.id'))
    name = Column(String(100))
    level = Column(String(20))
    pattern_name = Column(String(50))
    school = relationship("School", back_populates="departments")
    subjects = relationship("Subject", back_populates="department", cascade="all, delete-orphan")

class GridType(Base):
    __tablename__ = 'grid_type'
    id = Column(Integer, primary_key=True)
    name = Column(String(50))
    has_problem_column = Column(Integer)

class Subject(Base):
    __tablename__ = 'subject'
    id = Column(Integer, primary_key=True)
    dept_id = Column(Integer, ForeignKey('department.id'))
    name = Column(String(100))
    code = Column(String(20))
    semester = Column(Integer)
    pattern_name = Column(String(50))
    grid_type_id = Column(Integer, ForeignKey('grid_type.id'))
    department = relationship("Department", back_populates="subjects")
    grid_type = relationship("GridType")

class SubjectWeightage(Base):
    __tablename__ = "subject_weightage"

    id = Column(Integer, primary_key=True)
    subject_id = Column(Integer, ForeignKey("subject.id"), nullable=False)
    unit = Column(Integer, nullable=False)

    sec_a_count = Column(Integer, default=0)
    sec_b_count = Column(Integer, default=0)
    sec_c_count = Column(Integer, default=0)

    subject = relationship("Subject", backref="weightages")


class User(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True)
    username = Column(String(100), unique=True)
    password = Column(String(255))
    role = Column(String(50))
    school_access_id = Column(Integer, ForeignKey('school.id'), nullable=True)
    school = relationship("School")

class ExamPattern(Base):
    __tablename__ = 'exam_pattern'
    id = Column(Integer, primary_key=True)
    name = Column(String(50), unique=True, nullable=False)
    total_marks = Column(Integer, default=100)
    sections = relationship('PatternSection', back_populates='pattern', cascade="all, delete-orphan")

class PatternSection(Base):
    __tablename__ = 'pattern_section'
    id = Column(Integer, primary_key=True)
    pattern_id = Column(Integer, ForeignKey('exam_pattern.id'), nullable=False)
    section_name = Column(String(10), nullable=False)
    count = Column(Integer, nullable=False)
    total_in_paper = Column(Integer, nullable=False)
    marks = Column(Integer, nullable=False)
    note = Column(Text)
    pattern = relationship("ExamPattern", back_populates="sections")

Base.metadata.create_all(engine)

# --- HELPER FUNCTIONS ---
def get_recent_files():
    if not os.path.exists(UPLOAD_FOLDER): os.makedirs(UPLOAD_FOLDER)
    files = []
    for f in os.listdir(UPLOAD_FOLDER):
        if f.endswith('.docx'):
            path = os.path.join(UPLOAD_FOLDER, f)
            stats = os.stat(path)
            files.append({
                "File Name": f,
                "Created At": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M')
            })
    return sorted(files, key=lambda x: x['Created At'], reverse=True)

def get_accessible_schools(session, role=None, school_access_id=None):
    if role == "admin" or school_access_id is None:
        return session.query(School).all()
    else:
        school = session.get(School, school_access_id)
        return [school] if school else []

def safe_rerun():
    if hasattr(st, "rerun"): st.rerun()
    elif hasattr(st, "experimental_rerun"): st.experimental_rerun()

def get_exam_patterns(session):
    return session.query(ExamPattern).order_by(ExamPattern.name).all()

def load_question_bank(file):
    import pandas as pd
    import streamlit as st

    try:
        if file.name.endswith(".csv"):
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)

        # ✅ Normalize column names
        df.columns = (
            df.columns
            .astype(str)
            .str.strip()
            .str.lower()
            .str.replace(" ", "_")
        )

        # ✅ REQUIRED COLUMN MAPPING
        required_map = {
            "question": ["question"],
            "unit": ["unit"],
            "marks": ["marks"],
            "section": ["section"],
            "k_level": ["k_level", "klevel", "k_level_"]
        }

        col_map = {}
        for std, variants in required_map.items():
            for v in variants:
                if v in df.columns:
                    col_map[std] = v
                    break

        if "question" not in col_map or "unit" not in col_map:
            st.error(f"Missing required columns. Found: {list(df.columns)}")
            return None

        # ✅ Rename to STANDARD NAMES
        df = df.rename(columns={
            col_map["question"]: "question",
            col_map["unit"]: "unit",
            col_map.get("marks", ""): "marks",
            col_map.get("section", ""): "section",
            col_map.get("k_level", ""): "k_level"
        })

        st.session_state.q_bank = df
        return df

    except Exception as e:
        st.error(f"Failed to load file: {e}")
        return None
        return None

#def generate_question_paper(subject_name, questions_df):
def generate_question_paper(subject_name, questions_df, pattern_id, session):
    from docx import Document
    import re

    if questions_df is None or questions_df.empty:
        raise ValueError("Question bank is empty")

    # 🔹 Normalize column names
    questions_df.columns = [c.lower().strip() for c in questions_df.columns]

    # 🔹 Map Excel headers safely
    column_map = {
        "question": "question",
        "marks": "marks",
        "k level": "k_level",
        "k_level": "k_level",
        "section": "section"
    }

    questions_df.rename(columns=column_map, inplace=True)

    doc = Document()
    doc.add_heading(f"Question Paper – {subject_name}", level=1)

    # 🔹 Fetch exam pattern from DB
    pattern = session.get(ExamPattern, pattern_id)
    if not pattern:
        raise ValueError("Invalid exam pattern selected")

    for sec in pattern.sections:
        doc.add_heading(f"Section {sec.section_name}", level=2)
        doc.add_paragraph(
            f"Answer {sec.count} questions × {sec.marks} marks"
        )

        # 🔹 Filter questions matching marks (important)
        eligible = questions_df[
            questions_df["marks"].astype(str) == str(sec.marks)
        ]

        if eligible.empty:
            doc.add_paragraph("⚠ No matching questions available")
            continue

        selected = eligible.sample(
            min(sec.count, len(eligible)),
            random_state=None
        )

        for i, row in enumerate(selected.itertuples(), start=1):
            q_text = getattr(row, "question", "")
            klevel = getattr(row, "k_level", "")

            p = doc.add_paragraph(f"{i}. {q_text}")

            meta = []

            # Safe K-level handling
            if klevel:
                match = re.search(r"\d+", str(klevel))
                if match:
                    meta.append(f"K{match.group()}")

            meta.append(f"{sec.marks} Marks")

            if meta:
                doc.add_paragraph(" (" + " | ".join(meta) + ")")

    return doc

# --- CALLBACKS ---
def delete_school_callback(school_id):
    local_session = Session()
    try:
        dept_count = local_session.query(Department).filter(Department.school_id == school_id).count()
        if dept_count > 0:
            st.toast(f"⚠️ Cannot delete! School has {dept_count} departments.", icon="🚫")
        else:
            item = local_session.get(School, school_id)
            if item:
                local_session.delete(item)
                local_session.commit()
                st.toast(f"✅ Deleted School: {item.name}")
    except Exception as e: st.error(f"Error: {e}")
    finally: local_session.close()

def delete_department_callback(dept_id):
    local_session = Session()
    try:
        sub_count = local_session.query(Subject).filter(Subject.dept_id == dept_id).count()
        if sub_count > 0:
            st.toast(f"⚠️ Cannot delete! Department has {sub_count} subjects.", icon="🚫")
        else:
            item = local_session.get(Department, dept_id)
            if item:
                local_session.delete(item)
                local_session.commit()
                st.toast(f"✅ Deleted Department: {item.name}")
    except Exception as e: st.error(f"Error: {e}")
    finally: local_session.close()

def delete_subject_callback(subject_id):
    local_session = Session()
    try:
        item = local_session.get(Subject, subject_id)
        if item:
            local_session.delete(item)
            local_session.commit()
            st.toast(f"Deleted Subject: {item.name}")
    except Exception as e: st.error(f"Error: {e}")
    finally: local_session.close()

def delete_user_callback(user_id):
    local_session = Session()
    try:
        item = local_session.get(User, user_id)
        if item:
            local_session.delete(item)
            local_session.commit()
            st.toast(f"Deleted User: {item.username}")
    except Exception as e: st.error(f"Error: {e}")
    finally: local_session.close()

def delete_pattern_callback(pattern_id):
    local_session = Session()
    try:
        pat = local_session.get(ExamPattern, pattern_id)
        if pat:
            usage_count = local_session.query(Subject).filter(Subject.pattern_name == pat.name).count()
            if usage_count > 0:
                st.toast(f"⚠️ Cannot delete! Pattern used by {usage_count} subjects.", icon="🚫")
            else:
                local_session.delete(pat)
                local_session.commit()
                st.toast(f"✅ Deleted Pattern: {pat.name}")
    except Exception as e: st.error(f"Error: {e}")
    finally: local_session.close()

# --- LOGIN PAGE ---
def login_page():
    st.markdown(f"<h1 style='text-align:center'>{COLLEGE_NAME}</h1>", unsafe_allow_html=True)
    st.subheader("Admin Login")
    
    session = Session()
    admin_exists = False
    try:
        admin_exists = session.query(User).filter(User.role == 'admin').count() > 0
    except: pass
    finally: session.close()

    if not ADMIN_PASSWORD and not admin_exists:
        with st.form("create_admin"):
            new_user = st.text_input("Username", value=ADMIN_USERNAME)
            new_pass = st.text_input("Password", type="password")
            if st.form_submit_button("Create Admin"):
                if new_user and new_pass:
                    ls = Session()
                    ls.add(User(username=new_user, password=generate_password_hash(new_pass), role='admin'))
                    ls.commit(); ls.close()
                    st.success("Admin created!"); time.sleep(0.5); safe_rerun()
        return

    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        with st.form("login"):
            u = st.text_input("Username")
            p = st.text_input("Password", type="password")
            if st.form_submit_button("Login"):
                ls = Session()
                user = ls.query(User).filter(User.username==u).first()
                ls.close()
                valid = False
                if user:
                    try: valid = check_password_hash(user.password, p)
                    except: valid = (user.password == p)
                
                if not valid and u == ADMIN_USERNAME and p == ADMIN_PASSWORD:
                    st.session_state.update({'logged_in':True, 'user':u, 'role':'admin', 'school_access_id':None})
                    st.success("Logged in as Admin"); time.sleep(0.5); safe_rerun()
                    return

                if valid:
                    st.session_state.update({
                        'logged_in':True, 'user':user.username, 
                        'role':user.role, 'school_access_id':user.school_access_id
                    })
                    st.success(f"Welcome {user.role.title()}!"); time.sleep(0.5); safe_rerun()
                else:
                    st.error("Invalid credentials")

def sample_questions(df, unit, marks, count, section=None):
    import pandas as pd

    if df is None or df.empty:
        return []

    q = df[df["unit"] == unit]

    if section:
        q = q[q["section"].str.upper() == section.upper()]

    if "marks" in q.columns:
        q = q[q["marks"] == marks]

    if len(q) < count:
        return q.to_dict("records")

    return q.sample(count).to_dict("records")


# --- MAIN DASHBOARD ---
def main_dashboard():
    if not st.session_state.get('logged_in'):
        login_page(); return

    user_role = st.session_state.get('role', 'staff')
    is_admin = user_role == "admin"
    user_school_id = st.session_state.get('school_access_id')
    
    
    if 'q_bank' not in st.session_state:
        st.session_state.q_bank = pd.DataFrame()  # empty dataframe or load from DB

    # Now you can safely access it
    df = st.session_state.q_bank.copy()
    
    st.write(df)
    with st.sidebar:
        st.title("Admin Panel")
        if st.button("Logout"):
            st.session_state.clear(); safe_rerun()
        menu = st.radio("Menu", ["🏠 Home", "📄 Generate Paper", "🏫 Manage Schools", "🎓 Manage Departments", "📘 Manage Subjects", "📝 Manage Patterns", "👥 Manage Users", "📂 Downloads"])

    session = Session()

    
    # Only handle the Generate Paper workflow here
    if menu == "🏠 Home":
        if st.session_state.page == "SUBJECTS" and st.session_state.selected_department_id:
            if st.button("← Back to Departments"):
                st.session_state.page = "HOME"
                st.experimental_rerun()
            
            st.header("Select Subject & Upload Question Bank")
            
            # --- Subjects dropdown ---
            subjects = session.query(Subject).filter(
                Subject.dept_id == st.session_state.selected_department_id
            ).all()
            
            sel_subject = st.selectbox(
                "Select Subject",
                subjects,
                format_func=lambda x: x.name,
                key="qp_subject_select"
            )
            
            # --- Initialize q_bank safely ---
            if "q_bank" not in st.session_state:
                st.session_state.q_bank = None
            
            # --- File uploader ---
            uploaded_file = st.file_uploader(
                "Upload Question Bank (CSV/XLSX)",
                type=['csv', 'xlsx'],
                key="qp_file_upload"
            )
            
            if uploaded_file is not None:
                try:
                    if uploaded_file.name.endswith(".csv"):
                        st.session_state.q_bank = pd.read_csv(uploaded_file)
                    else:
                        st.session_state.q_bank = pd.read_excel(uploaded_file)
                    st.success("File uploaded successfully!")
                except Exception as e:
                    st.error(f"Error loading file: {e}")
            
            # --- Show uploaded Q-bank ---
            if st.session_state.q_bank is not None:
                df = st.session_state.q_bank.copy()
                df.columns = [c.strip() for c in df.columns]  # remove spaces
                st.dataframe(df)
            
            # --- Patterns ---
            patterns = []
            if sel_subject:
                patterns = session.query(ExamPattern).order_by(ExamPattern.name).all()
            
            if not patterns:
                st.warning("No exam patterns configured. Please add patterns first.")
                return
            
            pattern_map = {p.name: p.id for p in patterns}
            
            selected_pattern_name = st.selectbox(
                "Select Exam Pattern",
                list(pattern_map.keys()),
                key="qp_pattern_select"
            )
            
            # --- Generate Paper button ---
            generate_disabled = st.session_state.q_bank is None
            if st.button("Generate Question Paper", key="qp_generate_btn", disabled=generate_disabled):
                
                df = st.session_state.q_bank.copy()
                df.columns = [c.strip() for c in df.columns]
                
                # Get pattern
                pat = session.query(ExamPattern).filter_by(name=selected_pattern_name).first()
                if not pat:
                    st.error("No pattern assigned to this subject")
                    return
                
                # Normalize section names
                p_dict = {s.section_name.strip().replace(" ", "").upper(): s.marks for s in pat.sections}
                
                paper = []
                
                # Loop through Q-bank rows and generate questions
                for _, row in df.iterrows():
                    unit = row.get("Unit")
                    if not unit:
                        st.warning("Skipping row without Unit")
                        continue
                    
                    sec_a_cnt = int(row.get("Sec A", 0))
                    sec_b_cnt = int(row.get("Sec B", 0))
                    sec_c_cnt = int(row.get("Sec C", 0))
                    
                    if sec_a_cnt > 0 and "SECA" in p_dict:
                        paper.extend(sample_questions(st.session_state.q_bank, unit, p_dict["SECA"], sec_a_cnt, "A"))
                    if sec_b_cnt > 0 and "SECB" in p_dict:
                        paper.extend(sample_questions(st.session_state.q_bank, unit, p_dict["SECB"], sec_b_cnt, "B"))
                    if sec_c_cnt > 0 and "SECC" in p_dict:
                        paper.extend(sample_questions(st.session_state.q_bank, unit, p_dict["SECC"], sec_c_cnt, "C"))
                
                st.session_state.paper = paper
                st.session_state.gen_pat_id = pat.id
                
                if paper:
                    st.success(f"✅ Question paper generated ({len(paper)} Qs)")
                    
                    # --- Download button ---
                    from io import BytesIO
                    buffer = BytesIO()
                    paper.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        "Download Question Paper",
                        buffer,
                        file_name=f"{sel_subject.name}_QP.docx"
                    )
                else:
                    st.error("No questions could be generated. Check your Q-bank and pattern.")
#########################################################################################################
    try:
        # --- HOME ---
        if menu == "🏠 Home":
            if st.session_state.page == "HOME":
                st.header("Home — Visual Overview")
                schools = get_accessible_schools(session, user_role, user_school_id)
                c1, c2 = st.columns(2)
                lev = c1.selectbox("Filter Level:", ["All", "UG", "PG"])
                sch_opts = ["All"] + [s.name for s in schools]
                sch_filter = c2.selectbox("Filter School:", sch_opts)
                
                q = session.query(Department)
                if lev != "All": q = q.filter(Department.level == lev)
                if sch_filter != "All": q = q.filter(Department.school.has(School.name == sch_filter))
                else:
                    s_ids = [s.id for s in schools]
                    if s_ids: q = q.filter(Department.school_id.in_(s_ids))
                
                depts = q.all()
                if depts:
                    cols = st.columns(3)
                    for i, d in enumerate(depts):
                        with cols[i % 3]:
                            if st.button(
                                f"{d.name}\n{d.school.name if d.school else ''}\n{d.level}",
                                key=f"dept_{d.id}",
                                use_container_width=True
                            ):
                                st.session_state.selected_department_id = d.id
                                st.session_state.page = "SUBJECTS"
                                safe_rerun()
                else:
                    st.info("No departments found.")

            # --- DEPARTMENT SELECTED → SUBJECTS & GENERATE QP ---
            elif st.session_state.page == "SUBJECTS" and st.session_state.selected_department_id:
                if st.button("← Back to Departments"):
                    st.session_state.page = "HOME"
                    safe_rerun()

                st.header("Select Subject & Upload Question Bank")

                subjects = session.query(Subject).filter(
                    Subject.dept_id == st.session_state.selected_department_id
                ).all()

                sel_subject = st.selectbox(
                    "Select Subject",
                    subjects,
                    format_func=lambda x: x.name,
                    key="qp_subject_select"
                )

                uploaded_file = st.file_uploader(
                    "Upload Question Bank (CSV/XLSX)",
                    type=["csv", "xlsx"],
                    key="qp_file_upload"
                )
                
                uploaded_file = st.file_uploader("Upload your question bank CSV/Excel", type=['csv', 'xlsx'])
                if uploaded_file is not None:
                    if uploaded_file.name.endswith('.csv'):
                        st.session_state.q_bank = pd.read_csv(uploaded_file)
                    else:
                        st.session_state.q_bank = pd.read_excel(uploaded_file)
                    st.success("File uploaded successfully!")

    # Only show dataframe if q_bank exists
                if st.session_state.q_bank is not None:
                    df = st.session_state.q_bank.copy()
                    st.dataframe(df)
               # if uploaded_file:
                #    load_question_bank(uploaded_file)

                # ✅ ALWAYS define patterns
                patterns = []

                if "question_bank" in st.session_state and sel_subject:
                    patterns = session.query(ExamPattern).order_by(ExamPattern.name).all()

                if not patterns:
                    if "question_bank" in st.session_state:
                        st.warning("No exam patterns configured. Please add patterns first.")
                    return

                pattern_map = {p.name: p.id for p in patterns}

                selected_pattern_name = st.selectbox(
                    "Select Exam Pattern",
                    list(pattern_map.keys()),
                    key="qp_pattern_select"
                )

                if st.button("Generate Question Paper", key="qp_generate_btn"):
                    questions_df = st.session_state.question_bank

                    if isinstance(questions_df, dict):
                        questions_df = pd.DataFrame(questions_df)

                    paper = generate_question_paper(
                        sel_subject.name,
                        questions_df,
                        pattern_map[selected_pattern_name],
                        session
                    )

                    st.success("✅ Question Paper Generated")


                    from io import BytesIO
                    buffer = BytesIO()
                    paper.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        "Download Question Paper",
                        buffer,
                        file_name=f"{sel_subject.name}_QP.docx"
                    )

        # --- GENERATE PAPER (Direct Menu) ---
        elif menu == "📄 Generate Paper":
            if is_admin: st.warning("Please use management tools."); st.stop()
            st.header("Generate Paper")
            sl = get_accessible_schools(session, user_role, user_school_id)
            c1,c2 = st.columns(2)
            sch = c1.selectbox("School", sl, format_func=lambda x:x.name, key="gp_s")
            depts = session.query(Department).filter_by(school_id=sch.id).all() if sch else []
            dept = c2.selectbox("Department", depts, format_func=lambda x:x.name, key="gp_d")
            subs = session.query(Subject).filter_by(dept_id=dept.id).all() if dept else []
            sub = st.selectbox("Subject", subs, format_func=lambda x:x.name, key="gp_sub")
            
            if sub:
                # --- WEIGHTAGE LOGIC ADDED HERE ---
                #w_list = session.query(SubjectWeightage).filter_by(subject_id=sub.id).order_by(SubjectWeightage.unit).all()
                w_list = session.query(SubjectWeightage)\
    .filter_by(subject_id=sub.id)\
    .order_by(SubjectWeightage.unit)\
    .all()

                if not w_list:
                    for u in range(1, 6): session.add(SubjectWeightage(subject_id=sub.id, unit=u)); 
                    session.commit(); w_list = session.query(SubjectWeightage).filter_by(subject_id=sub.id).order_by(SubjectWeightage.unit).all()
                
                df_w = pd.DataFrame([{"Unit": w.unit, "Sec A": w.sec_a_count, "Sec B": w.sec_b_count, "Sec C": w.sec_c_count} for w in w_list])
                
                g1, g2 = st.columns([1, 1])
                with g1:
                    st.markdown("### 1. Weightage")
                    # Editable dataframe for weightage
                    edited_df = st.data_editor(
                        df_w, 
                        hide_index=True, 
                        use_container_width=True, 
                        key="gp_w_edit",
                        column_config={
                            "Unit": st.column_config.NumberColumn(disabled=True),
                            "Sec A": st.column_config.NumberColumn(min_value=0, max_value=10),
                            "Sec B": st.column_config.NumberColumn(min_value=0, max_value=10),
                            "Sec C": st.column_config.NumberColumn(min_value=0, max_value=10)
                        }
                    )
                if st.button(
    "Generate Paper",
    key="gp_btn",
    disabled=("q_bank" not in st.session_state),
    type="primary"
):

    # ✅ MUST BE FIRST LINE
                    p_dict = {}

            pat = session.query(ExamPattern).filter_by(name=sub.pattern_name).first()
            if not pat:
                st.error("No pattern assigned to this subject")
                st.stop()

    # ✅ Normalize section names
            p_dict = {s.section_name.strip().replace(" ", "").upper(): s.marks for s in pat.sections}

    # ✅ Ensure uploaded df exists
            df = st.session_state.q_bank.copy()
            df.columns = [c.strip() for c in df.columns]  # remove extra spaces

            paper = []

            for _, row in df.iterrows():
                unit = row.get("Unit")
                if not unit:
                    st.warning("Skipping row without Unit")
                    continue

        # Get counts safely
                sec_a_cnt = int(row.get("Sec A", 0))
                sec_b_cnt = int(row.get("Sec B", 0))
                sec_c_cnt = int(row.get("Sec C", 0))

                if sec_a_cnt > 0 and "SECA" in p_dict:
                    paper.extend(sample_questions(st.session_state.q_bank, unit, p_dict["SECA"], sec_a_cnt, "A"))
                if sec_b_cnt > 0 and "SECB" in p_dict:
                    paper.extend(sample_questions(st.session_state.q_bank, unit, p_dict["SECB"], sec_b_cnt, "B"))
                if sec_c_cnt > 0 and "SECC" in p_dict:
                    paper.extend(sample_questions(st.session_state.q_bank, unit, p_dict["SECC"], sec_c_cnt, "C"))

    # ✅ Assign once after loop
            st.session_state.paper = paper
            st.session_state.gen_pat_id = pat.id

            if paper:
                st.success(f"Question paper generated ({len(paper)} Qs)")
                st.rerun()
            else:
                st.error("No questions could be generated. Check your Q-bank and pattern.")


        # --- SCHOOLS ---
        elif menu == "🏫 Manage Schools":
            if user_school_id: st.warning("Restricted"); st.stop()
            st.header("Manage Schools")
            c1, c2 = st.columns([1, 2])
            with c1:
                with st.form("add_s"):
                    n = st.text_input("Name")
                    if st.form_submit_button("Add") and n:
                        ls = Session()
                        if not ls.query(School).filter_by(name=n).first():
                            ls.add(School(name=n)); ls.commit(); st.success("Added"); time.sleep(0.5); safe_rerun()
                        ls.close()
            with c2:
                lst = session.query(School).all()
                if lst:
                    st.dataframe(pd.DataFrame([{"ID":s.id, "Name":s.name} for s in lst]), use_container_width=True)
                    sel = st.selectbox("Delete", lst, format_func=lambda x:x.name, key="del_s_sel")
                    st.button("Delete", on_click=delete_school_callback, args=(sel.id,))

        # --- DEPARTMENTS ---
        elif menu == "🎓 Manage Departments":
            st.header("Manage Departments")
            sl = get_accessible_schools(session, user_role, user_school_id)
            c1, c2 = st.columns([1, 2])
            with c1:
                with st.form("add_d"):
                    n = st.text_input("Name"); l = st.selectbox("Level", ["UG", "PG"]); s = st.selectbox("School", sl, format_func=lambda x:x.name)
                    if st.form_submit_button("Add") and n and s:
                        ls = Session(); ls.add(Department(name=n, level=l, school_id=s.id)); ls.commit(); st.success("Added"); time.sleep(0.5); safe_rerun()
            with c2:
                q = session.query(Department)
                if user_school_id: q = q.filter_by(school_id=user_school_id)
                res = q.all()
                if res:
                    st.dataframe(pd.DataFrame([{"Name":d.name, "Level":d.level, "School":d.school.name} for d in res]), use_container_width=True)
                    sel = st.selectbox("Delete", res, format_func=lambda x:f"{x.name} ({x.level})", key="del_d_sel")
                    st.button("Delete", on_click=delete_department_callback, args=(sel.id,))

        # --- SUBJECTS ---
        elif menu == "📘 Manage Subjects":
            st.header("Manage Subjects")
            c1, c2 = st.columns([1, 2])
            depts = session.query(Department).all()
            if user_school_id: depts = [d for d in depts if d.school_id == user_school_id]
            pats = session.query(ExamPattern).all()
            grids = session.query(GridType).all()

            with c1:
                with st.form("add_sub"):
                    n = st.text_input("Name"); c = st.text_input("Code"); sem = st.number_input("Sem", 1, 6)
                    d = st.selectbox("Dept", depts, format_func=lambda x:x.name) if depts else None
                    p = st.selectbox("Pattern", pats, format_func=lambda x:x.name) if pats else None
                    g = st.selectbox("Grid", grids, format_func=lambda x:x.name) if grids else None
                    
                    if st.form_submit_button("Add") and n and c and d and p and g:
                        ls = Session()
                        ls.add(Subject(name=n, code=c, semester=sem, dept_id=d.id, pattern_name=p.name, grid_type_id=g.id))
                        ls.commit(); st.success("Added"); time.sleep(0.5); safe_rerun(); ls.close()

            with c2:
                q = session.query(Subject)
                if user_school_id: 
                    d_ids = [d.id for d in depts]
                    q = q.filter(Subject.dept_id.in_(d_ids)) if d_ids else q.filter(False)
                res = q.all()
                if res:
                    st.dataframe(pd.DataFrame([{"Code":x.code, "Name":x.name, "Dept":x.department.name} for x in res]), use_container_width=True)
                    sel = st.selectbox("Delete", res, format_func=lambda x:x.code, key="del_sub_sel")
                    st.button("Delete", on_click=delete_subject_callback, args=(sel.id,))

        # --- PATTERNS ---
        elif menu == "📝 Manage Patterns":
            if user_school_id: st.warning("Restricted"); st.stop()
            st.header("Manage Exam Patterns")
            c1, c2 = st.columns([1, 2])
            with c1:
                with st.form("add_pat"):
                    pn = st.text_input("Name"); pm = st.number_input("Max Marks", 50, 100, 100)
                    secs = []
                    for s in ['SecA', 'SecB', 'SecC']:
                        st.caption(s)
                        c = st.columns(3)
                        cnt = c[0].number_input("Cnt", 0, 10, key=f"c{s}")
                        mks = c[1].number_input("Mk", 0, 20, key=f"m{s}")
                        tot = c[2].number_input("Tot", 0, 20, key=f"t{s}")
                        if cnt: secs.append((s, cnt, mks, tot))
                    if st.form_submit_button("Create") and pn:
                        ls = Session()
                        if not ls.query(ExamPattern).filter_by(name=pn).first():
                            p = ExamPattern(name=pn, total_marks=pm); ls.add(p); ls.flush()
                            for sn, sc, sm, stt in secs:
                                ls.add(PatternSection(pattern_id=p.id, section_name=sn, count=sc, marks=sm, total_in_paper=stt))
                            ls.commit(); st.success("Added"); time.sleep(0.5); safe_rerun()
                        ls.close()

            with c2:
                res = session.query(ExamPattern).all()
                if res:
                    # Enhanced Data Table for Patterns
                    table_data = []
                    for p in res:
                        secs = {s.section_name: s.count for s in p.sections}
                        sub_count = session.query(Subject).filter(Subject.pattern_name == p.name).count()
                        table_data.append({
                            "Name": p.name,
                            "Marks": p.total_marks,
                            "Unit": 5, 
                            "Sec A": secs.get("SecA", 0),
                            "Sec B": secs.get("SecB", 0),
                            "Sec C": secs.get("SecC", 0),
                            "Subject Weightage": sub_count
                        })
                    st.dataframe(pd.DataFrame(table_data), use_container_width=True)
                    sel = st.selectbox("Delete", res, format_func=lambda x:x.name, key="del_pat_sel")
                    st.button("Delete", on_click=delete_pattern_callback, args=(sel.id,))

        # --- USERS ---
        elif menu == "👥 Manage Users":
            if user_school_id: st.warning("Access Denied"); st.stop()
            st.header("Manage Users")
            c1, c2 = st.columns([1, 2])
            with c1:
                with st.form("add_u"):
                    u = st.text_input("User"); p = st.text_input("Pass")
                    r = st.selectbox("Role", ["staff", "admin"])
                    s = st.selectbox("School Access", [None]+session.query(School).all(), format_func=lambda x:x.name if x else "All")
                    if st.form_submit_button("Add") and u and p:
                        ls = Session()
                        ls.add(User(username=u, password=generate_password_hash(p), role=r, school_access_id=s.id if s else None))
                        ls.commit(); st.success("Added"); time.sleep(0.5); safe_rerun(); ls.close()
            with c2:
                res = session.query(User).all()
                st.dataframe(pd.DataFrame([{"User":u.username, "Role":u.role} for u in res]), use_container_width=True)
                if res:
                    sel = st.selectbox("Delete", res, format_func=lambda x:x.username, key="del_u_sel")
                    st.button("Delete", on_click=delete_user_callback, args=(sel.id,))

        # --- DOWNLOADS ---
        elif menu == "📂 Downloads":
            st.header("Files")
            files = get_recent_files()
            if files:
                for f in files:
                    c1, c2 = st.columns([4, 1])
                    c1.write(f"{f['File Name']} ({f['Created At']})")
                    with open(os.path.join(UPLOAD_FOLDER, f['File Name']), "rb") as fl:
                        c2.download_button("Download", fl, file_name=f['File Name'], key=f"dl_{f['File Name']}")
                    st.divider()
            else: st.info("No files generated yet.")

    finally: session.close()

def build_question_paper_docx(questions, subject_name, subject_id, session):
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pandas as pd

    if not questions:
        raise ValueError("No questions to generate paper")

    subject_obj = session.get(Subject, subject_id) if subject_id else None
    pattern_key = subject_obj.pattern_name if subject_obj else "Pattern_1"
    pattern_data = get_pattern_dict(pattern_key)

    df_paper = pd.DataFrame(questions)

    doc = Document()

    # --- Page Margins ---
    sec = doc.sections[0]
    sec.top_margin = Cm(1.27)
    sec.bottom_margin = Cm(1.27)
    sec.right_margin = Cm(1.27)
    sec.left_margin = Cm(1.5)

    # --- REG NO ---
    reg_p = doc.add_paragraph()
    reg_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    reg_p.paragraph_format.space_after = Pt(0)
    reg_p.add_run("REG. NO : _______________________").bold = True

    # --- Subject Title ---
    title = doc.add_paragraph(f"\n{subject_name}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    # --- Group by Section ---
    if "section" in df_paper.columns:
        grouped = df_paper.groupby("section")
    else:
        grouped = [("A", df_paper)]

    q_no = 1
    for sec_name, gdf in grouped:
        doc.add_paragraph(f"\nSECTION {sec_name}").runs[0].bold = True

        for _, row in gdf.iterrows():
            p = doc.add_paragraph(f"{q_no}. {row['question']}")
            p.paragraph_format.space_after = Pt(6)
            q_no += 1

    return doc


if __name__ == "__main__":
    if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
    main_dashboard()